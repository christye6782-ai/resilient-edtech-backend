"""Computer-vision pipeline that turns an uploaded lesson plan into text.

A rural teacher often only has a *photo* or *scan* of a handwritten or printed
lesson plan. This module is the CV model the backend connects to:

    raw bytes  ->  OpenCV preprocessing (denoise, deskew, adaptive threshold)
               ->  Tesseract OCR  ->  extracted text + confidence + preview

It also accepts PDF / DOCX / plain-text uploads and extracts their text directly,
so the same endpoint handles every format a teacher might have on hand.
"""
from __future__ import annotations

import base64
import io
import logging

from .config import settings
from .schemas import ExtractionResult

logger = logging.getLogger("resilient-edtech.cv")

IMAGE_EXT = {".png", ".jpg", ".jpeg", ".bmp", ".tif", ".tiff", ".webp"}


# --------------------------------------------------------------------------- #
# Public entry point
# --------------------------------------------------------------------------- #

def extract_from_upload(filename: str, content: bytes) -> ExtractionResult:
    """Route an upload to the right extractor based on its extension."""
    name = (filename or "").lower()
    ext = name[name.rfind("."):] if "." in name else ""

    if ext in IMAGE_EXT:
        return _ocr_image(content)
    if ext == ".pdf":
        return _extract_pdf(content)
    if ext == ".docx":
        return _extract_docx(content)
    # default: treat as plain text
    return _extract_text(content)


# --------------------------------------------------------------------------- #
# Computer-vision OCR for images
# --------------------------------------------------------------------------- #

def _ocr_image(content: bytes) -> ExtractionResult:
    result = ExtractionResult(source_type="image", method="OCR (OpenCV + Tesseract)")
    try:
        import cv2
        import numpy as np
    except Exception:  # noqa: BLE001
        result.warnings.append("OpenCV/numpy not available — cannot process images.")
        return result

    # Decode the uploaded bytes into an OpenCV image (BGR).
    arr = np.frombuffer(content, dtype=np.uint8)
    img = cv2.imdecode(arr, cv2.IMREAD_COLOR)
    if img is None:
        result.warnings.append("Could not decode the image file.")
        return result

    processed = _preprocess(cv2, np, img)

    # Encode the cleaned-up page so the UI can show what the CV model 'saw'.
    ok, buf = cv2.imencode(".png", processed)
    if ok:
        result.preview_image = "data:image/png;base64," + base64.b64encode(buf.tobytes()).decode("ascii")

    text, confidence = _run_tesseract(processed)
    if text is None:
        result.warnings.append(
            "Tesseract OCR engine not found. Install it (conda install -c conda-forge "
            "tesseract, or the UB-Mannheim Windows build) and set TESSERACT_CMD in .env. "
            "You can also paste the lesson text directly."
        )
        return result

    result.text = text.strip()
    result.confidence = confidence
    result.word_count = len(result.text.split())
    if result.word_count == 0:
        result.warnings.append("OCR ran but found no readable text — try a clearer, higher-contrast photo.")
    return result


def _preprocess(cv2, np, img):
    """Classic document-scan CV pipeline: grayscale -> denoise -> deskew -> threshold."""
    gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)

    # Upscale small images so thin strokes survive thresholding.
    h, w = gray.shape[:2]
    if max(h, w) < 1000:
        scale = 1000.0 / max(h, w)
        gray = cv2.resize(gray, None, fx=scale, fy=scale, interpolation=cv2.INTER_CUBIC)

    gray = cv2.fastNlMeansDenoising(gray, h=15)
    gray = _deskew(cv2, np, gray)

    # Adaptive threshold copes with uneven lighting in classroom photos.
    binary = cv2.adaptiveThreshold(
        gray, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C, cv2.THRESH_BINARY, 31, 15
    )
    return binary


def _deskew(cv2, np, gray):
    """Rotate the page so text lines are horizontal (improves OCR a lot)."""
    inverted = cv2.bitwise_not(gray)
    coords = np.column_stack(np.where(inverted > 0))
    if coords.size == 0:
        return gray
    angle = cv2.minAreaRect(coords)[-1]
    if angle < -45:
        angle = 90 + angle
    if abs(angle) < 0.5:
        return gray  # already straight
    (h, w) = gray.shape[:2]
    matrix = cv2.getRotationMatrix2D((w // 2, h // 2), angle, 1.0)
    return cv2.warpAffine(
        gray, matrix, (w, h), flags=cv2.INTER_CUBIC, borderMode=cv2.BORDER_REPLICATE
    )


def _run_tesseract(processed):
    """Return (text, mean_confidence) or (None, None) if Tesseract is missing."""
    try:
        import pytesseract
        from PIL import Image
    except Exception:  # noqa: BLE001
        return None, None

    if settings.tesseract_cmd:
        pytesseract.pytesseract.tesseract_cmd = settings.tesseract_cmd

    try:
        pil = Image.fromarray(processed)
        data = pytesseract.image_to_data(pil, output_type=pytesseract.Output.DICT)
    except Exception as exc:  # noqa: BLE001 — usually 'tesseract is not installed'
        logger.warning("Tesseract failed: %s", exc)
        return None, None

    words, confs = [], []
    for token, conf in zip(data["text"], data["conf"]):
        token = token.strip()
        try:
            conf_val = float(conf)
        except (TypeError, ValueError):
            conf_val = -1.0
        if token and conf_val >= 0:
            words.append(token)
            confs.append(conf_val)

    text = " ".join(words)
    mean_conf = round(sum(confs) / len(confs), 1) if confs else 0.0
    return text, mean_conf


# --------------------------------------------------------------------------- #
# Document extractors (PDF / DOCX / text)
# --------------------------------------------------------------------------- #

def _extract_pdf(content: bytes) -> ExtractionResult:
    result = ExtractionResult(source_type="pdf", method="PDF text extraction (pypdf)")
    try:
        from pypdf import PdfReader
    except Exception:  # noqa: BLE001
        result.warnings.append("pypdf not available — cannot read PDFs.")
        return result
    try:
        reader = PdfReader(io.BytesIO(content))
        pages = [(page.extract_text() or "") for page in reader.pages]
        result.text = "\n".join(pages).strip()
    except Exception as exc:  # noqa: BLE001
        result.warnings.append(f"Could not read PDF: {exc}")
        return result

    result.word_count = len(result.text.split())
    if result.word_count == 0:
        result.warnings.append(
            "This PDF appears to be a scanned image with no embedded text. "
            "Export a page as JPG/PNG and upload that so the OCR pipeline can read it."
        )
    return result


def _extract_docx(content: bytes) -> ExtractionResult:
    result = ExtractionResult(source_type="docx", method="DOCX text extraction (python-docx)")
    try:
        from docx import Document
    except Exception:  # noqa: BLE001
        result.warnings.append("python-docx not available — cannot read .docx files.")
        return result
    try:
        doc = Document(io.BytesIO(content))
        parts = [p.text for p in doc.paragraphs]
        for table in doc.tables:
            for row in table.rows:
                parts.append(" | ".join(cell.text for cell in row.cells))
        result.text = "\n".join(t for t in parts if t.strip()).strip()
    except Exception as exc:  # noqa: BLE001
        result.warnings.append(f"Could not read DOCX: {exc}")
        return result
    result.word_count = len(result.text.split())
    return result


def _extract_text(content: bytes) -> ExtractionResult:
    result = ExtractionResult(source_type="text", method="Direct text decode")
    for encoding in ("utf-8", "latin-1"):
        try:
            result.text = content.decode(encoding).strip()
            break
        except UnicodeDecodeError:
            continue
    result.word_count = len(result.text.split())
    if not result.text:
        result.warnings.append("Could not decode this file as text. Supported: image, PDF, DOCX, TXT.")
    return result
